"""
Attachment parsing service for extracting readable text from various file formats.

Phase 1: Supports PDF, DOCX, TXT, and JSON files.
Extracts text for read-only contextual input to requirements agent.
"""
import json
import mimetypes
from typing import Optional
from pathlib import Path


class AttachmentParserError(Exception):
    """Error during attachment parsing."""
    pass


def extract_text_from_attachment(
    file_content: bytes,
    filename: str,
    mime_type: Optional[str] = None
) -> str:
    """
    Extract readable text from an attachment.
    
    Supported formats:
    - PDF (.pdf)
    - DOCX (.docx)
    - TXT (.txt)
    - JSON (.json)
    
    Args:
        file_content: Raw file content as bytes
        filename: Original filename (used for format detection)
        mime_type: Optional MIME type (if not provided, inferred from filename)
        
    Returns:
        Extracted text as string
        
    Raises:
        AttachmentParserError: If file format is unsupported or parsing fails
    """
    # Infer MIME type if not provided
    if not mime_type:
        mime_type, _ = mimetypes.guess_type(filename)
    
    # Determine file type from extension or MIME type
    file_ext = Path(filename).suffix.lower()
    
    try:
        if file_ext == '.txt' or mime_type == 'text/plain':
            return _extract_text_from_txt(file_content)
        elif file_ext == '.json' or mime_type == 'application/json':
            return _extract_text_from_json(file_content)
        elif file_ext == '.pdf' or mime_type == 'application/pdf':
            return _extract_text_from_pdf(file_content)
        elif file_ext == '.docx' or mime_type == 'application/vnd.openxmlformats-officedocument.wordprocessingml.document':
            return _extract_text_from_docx(file_content)
        else:
            raise AttachmentParserError(
                f"Unsupported file format: {filename} (MIME type: {mime_type}). "
                f"Supported formats: PDF, DOCX, TXT, JSON"
            )
    except Exception as e:
        if isinstance(e, AttachmentParserError):
            raise
        raise AttachmentParserError(f"Failed to parse attachment {filename}: {str(e)}")


def _extract_text_from_txt(file_content: bytes) -> str:
    """Extract text from plain text file."""
    try:
        # Try UTF-8 first
        return file_content.decode('utf-8')
    except UnicodeDecodeError:
        # Fallback to latin-1 (covers most cases)
        return file_content.decode('latin-1', errors='ignore')


def _extract_text_from_json(file_content: bytes) -> str:
    """Extract text from JSON file."""
    try:
        text = file_content.decode('utf-8')
        # Parse JSON to ensure it's valid, then return formatted string
        data = json.loads(text)
        # Return pretty-printed JSON as text
        return json.dumps(data, indent=2, ensure_ascii=False)
    except json.JSONDecodeError as e:
        # If JSON is invalid, return raw text with error note
        return f"[Invalid JSON - raw content]\n{file_content.decode('utf-8', errors='ignore')}"


def _extract_text_from_pdf(file_content: bytes) -> str:
    """
    Extract text from PDF file.
    
    Note: This is a basic implementation. For production, consider using
    libraries like PyPDF2, pdfplumber, or pymupdf.
    """
    try:
        # Try to import PyPDF2 (common PDF library)
        try:
            import PyPDF2
            from io import BytesIO
            
            pdf_file = BytesIO(file_content)
            pdf_reader = PyPDF2.PdfReader(pdf_file)
            
            text_parts = []
            for page_num, page in enumerate(pdf_reader.pages, start=1):
                try:
                    page_text = page.extract_text()
                    if page_text.strip():
                        text_parts.append(f"[Page {page_num}]\n{page_text}")
                except Exception as e:
                    text_parts.append(f"[Page {page_num} - extraction failed: {str(e)}]")
            
            return "\n\n".join(text_parts) if text_parts else "[PDF file - no extractable text]"
        except ImportError:
            # PyPDF2 not available - return placeholder
            return "[PDF file - text extraction requires PyPDF2 library. Install with: pip install PyPDF2]"
    except Exception as e:
        raise AttachmentParserError(f"PDF extraction failed: {str(e)}")


def _extract_text_from_docx(file_content: bytes) -> str:
    """
    Extract text from DOCX file.
    
    Note: This requires python-docx library.
    """
    try:
        try:
            from docx import Document
            from io import BytesIO
            
            docx_file = BytesIO(file_content)
            doc = Document(docx_file)
            
            text_parts = []
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_parts.append(paragraph.text)
            
            # Also extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
                    if row_text:
                        text_parts.append(row_text)
            
            return "\n".join(text_parts) if text_parts else "[DOCX file - no extractable text]"
        except ImportError:
            # python-docx not available - return placeholder
            return "[DOCX file - text extraction requires python-docx library. Install with: pip install python-docx]"
    except Exception as e:
        raise AttachmentParserError(f"DOCX extraction failed: {str(e)}")

